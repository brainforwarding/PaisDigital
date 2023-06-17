import os
import docx
import pandas as pd
import openai
from concurrent.futures import ThreadPoolExecutor


# Set the OpenAI API key
openai.api_key = "TU API KEY"

def read_essay(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def analyze(essay, rubric):
    conversation_history = [
        {"role": "system", "content": 
         """
Ayúdame a corregir el ensayo que te daré en base a una rúbrica que también te entregaré.

Siempre es importante tratar a los estudiantes con amabilidad y mantener altas expectativas hacia ellos. Además, utliza emojies cuando te dirijas a ellos.

Cuando analices un ensayo, debes incluir el nombre del estudiante, un resumen del ensayo, un análisis cualitativo del ensayo según cada criterio de la rúbrica, y una calificación numérica basada en la rúbrica dada.

Además, entrega a los estudiantes una carta con una retroalimentación personalizada sobre su ensayo. Incluye un saludo y un breve comentario introductorio, así como al menos dos aspectos destacables del trabajo y al menos dos áreas de mejora. Finalmente, cierra con un mensaje positivo.
         """},
                 {"role": "user", "content": 
         """
Ejemplo de analisis de un ensayo según el formato de salida para una rúbrica dada:

Nombre del estudiante: Antonia Lizana

Resumen del Ensayo: El ensayo discute sobre la biodiversidad marina, su importancia para el planeta y la humanidad, y los problemas de contaminación que la afectan. La autora sugiere algunas acciones para mitigar el daño.

Análisis:
- Contenido: Antonia tiene un enfoque relevante al hablar sobre la biodiversidad marina y la contaminación. Se expone una tesis clara y se argumenta en torno a la importancia de la biodiversidad marina y los efectos negativos de la contaminación. Sin embargo, faltan ejemplos y evidencias más concretas que respalden sus argumentos. No se utiliza pregunta retórica de manera efectiva.
- Redacción: La redacción es clara en general, pero hay ciertas frases que resultan confusas, y la estructura de las oraciones y párrafos podría variar más. La conclusión podría ser más sólida.
- Ortografía: Hay algunos errores ortográficos y de puntuación en el texto.
- Creatividad: Antonia muestra un enfoque interesante para el tema elegido, aunque podría haber incorporado ideas más creativas o frescas en su ensayo.

Calificación: 
- Contenido: 6
- Redacción: 3
- Ortografía: 2
- Creatividad: 3
Total: 14

Retroalimentación para la estudiante:

Hola Antonia, 😊

Has hecho un buen trabajo al abordar un tema tan relevante como es la biodiversidad marina y la contaminación. 👏 Me gusta tu enfoque en la importancia de esta biodiversidad y cómo nuestra acción puede influir en ella.

Aspectos destacables:
- Tu claridad para definir la biodiversidad marina y por qué es importante.
- Tu capacidad para conectar el tema con problemas más amplios como la contaminación y su impacto en la humanidad y la vida marina.

Áreas de mejora:
- Te sugiero añadir más ejemplos y evidencias para reforzar tus argumentos. 📚
- Trabaja en la estructura de tus frases y párrafos para mejorar la fluidez y claridad de tu ensayo. Y no olvides revisar la ortografía y puntuación.

Recuerda, la práctica hace al maestro. Sigue escribiendo y explorando estos temas tan importantes. Estoy segura de que tu próximo ensayo será aún mejor. 🌟

¡Sigue así!
         """},
        {"role": "user", "content": f"Aquí tienes la rúbrica de calificación: \n\n{rubric}"},
        {"role": "user", "content": f"Aquí tienes el ensayo para analizar: \n\n{essay}"}

    ]
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=conversation_history,
    )
    print(response.choices[0].finish_reason)
    return response.choices[0].message.content.strip()

def write_feedback(doc, analysis):
    doc.add_paragraph(f"\n{analysis}")
    doc.add_paragraph("\n---------------------------------------------\n")

def grade_essays(essays_folder, rubrica):
    # Load the rubric
    rubric_df = pd.read_csv(rubrica)
    rubric = rubric_df.to_string(index=False)

    # Iterate over all the docx files in the essays folder
    with ThreadPoolExecutor() as executor:
        for file in os.listdir(essays_folder):
            if file.endswith('.docx'):
                file_path = os.path.join(essays_folder, file)
                # Submit the task to the executor
                executor.submit(analyze_and_write_feedback, file_path, rubric)

def analyze_and_write_feedback(file_path, rubric):
    # Read the essay
    essay = read_essay(file_path)
    # Analyze
    analysis = analyze(essay, rubric)
    # Prepare the output document
    doc = docx.Document()
    write_feedback(doc, analysis)
    # Save the output document with a new name
    new_file_path = os.path.splitext(file_path)[0] + "_analysis.docx"
    doc.save(new_file_path)

if __name__ == "__main__":
    grade_essays("essays_folder", "rubrica.csv")
